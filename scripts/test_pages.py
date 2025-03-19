#!/usr/bin/env python3
import sys
import json
import os
from docxtpl import DocxTemplate
import logging
from datetime import datetime
from docx import Document
import requests
from pathlib import Path

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def count_pdf_pages(docx_path):
    """Конвертирует DOCX в PDF и считает страницы."""
    try:
        # URL для Gotenberg из переменной окружения или по умолчанию
        gotenberg_url = os.getenv('GOTENBERG_URL', 'http://localhost:3000')
        endpoint = f"{gotenberg_url}/forms/libreoffice/convert"

        logger.info(f"Converting {docx_path} to PDF using Gotenberg at {endpoint}")
        
        # Готовим файл для отправки
        with open(docx_path, 'rb') as f:
            files = {'file': (os.path.basename(docx_path), f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            
            # Отправляем запрос в Gotenberg
            response = requests.post(endpoint, files=files)
            
            if response.status_code == 200:
                # Сохраняем PDF во временный файл
                pdf_path = str(Path(docx_path).with_suffix('.pdf'))
                with open(pdf_path, 'wb') as f:
                    f.write(response.content)
                
                # Читаем количество страниц из PDF
                import PyPDF2
                with open(pdf_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    num_pages = len(pdf.pages)
                    logger.info(f"PDF has {num_pages} pages")
                
                # Удаляем временный PDF файл
                os.remove(pdf_path)
                return num_pages
            else:
                logger.error(f"Gotenberg conversion failed: {response.status_code}")
                return None
    except Exception as e:
        logger.error(f"Error converting to PDF: {e}", exc_info=True)
        return None

def get_page_number_from_footer(doc):
    """Получает номер страницы из нижнего колонтитула последней секции."""
    try:
        # Получаем последнюю секцию
        section = doc.sections[-1]
        
        # Получаем нижний колонтитул
        footer = section.footer
        
        logger.info("Analyzing footer content:")
        for paragraph in footer.paragraphs:
            text = paragraph.text.strip()
            logger.info(f"Footer text: '{text}'")
            
            # Ищем число в тексте
            import re
            numbers = re.findall(r'\d+', text)
            if numbers:
                # Берем последнее найденное число
                page_number = int(numbers[-1])
                logger.info(f"Found page number in footer: {page_number}")
                return page_number
                
        logger.warning("No page number found in footer")
        return None
    except Exception as e:
        logger.error(f"Error getting page number from footer: {e}", exc_info=True)
        return None

def get_docx_pages(docx_path):
    """Получает количество страниц из DOCX файла."""
    try:
        logger.info(f"Reading document from: {docx_path}")
        doc = Document(docx_path)
        
        # Пробуем получить номер страницы из колонтитула
        page_count = get_page_number_from_footer(doc)
        if page_count is not None:
            logger.info(f"Using page number from footer: {page_count}")
            return page_count
            
        # Если не удалось получить из колонтитула, пробуем через PDF
        pdf_pages = count_pdf_pages(docx_path)
        if pdf_pages is not None:
            return pdf_pages
            
        # Если не удалось через PDF, используем резервный метод подсчета разрывов
        # Подсчитываем количество разрывов страниц
        page_breaks = 1  # Начинаем с 1, так как первая страница не имеет разрыва
        
        # Подсчет в параграфах
        paragraph_breaks = 0
        for i, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                breaks = len(run._element.xpath('.//w:br[@w:type="page"]'))
                if breaks > 0:
                    paragraph_breaks += breaks
                    logger.info(f"Found {breaks} page break(s) in paragraph {i}")
        
        logger.info(f"Total page breaks in paragraphs: {paragraph_breaks}")
        page_breaks += paragraph_breaks
        
        # Подсчет в таблицах
        table_breaks = 0
        for i, table in enumerate(doc.tables):
            table_cell_breaks = 0
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            breaks = len(run._element.xpath('.//w:br[@w:type="page"]'))
                            if breaks > 0:
                                table_cell_breaks += breaks
                                logger.info(f"Found {breaks} page break(s) in table {i}")
            table_breaks += table_cell_breaks
        
        logger.info(f"Total page breaks in tables: {table_breaks}")
        page_breaks += table_breaks
        
        # Проверяем секции документа
        sections = len(doc.sections)
        logger.info(f"Document has {sections} section(s)")
        
        # Дополнительная проверка через свойства документа
        try:
            core_props = doc.core_properties
            if hasattr(core_props, 'page_count') and core_props.page_count:
                logger.info(f"Document core properties report {core_props.page_count} pages")
        except Exception as e:
            logger.warning(f"Could not get page count from core properties: {e}")

        # Проверяем другие возможные индикаторы страниц
        logger.info("Checking additional page indicators:")
        logger.info(f"Total paragraphs: {len(doc.paragraphs)}")
        logger.info(f"Total tables: {len(doc.tables)}")
        logger.info(f"Document sections: {len(doc.sections)}")
        
        # Анализ секций
        for i, section in enumerate(doc.sections):
            logger.info(f"Section {i+1}:")
            logger.info(f"  - Page height: {section.page_height}")
            logger.info(f"  - Page width: {section.page_width}")
            logger.info(f"  - Top margin: {section.top_margin}")
            logger.info(f"  - Bottom margin: {section.bottom_margin}")
            logger.info(f"  - Left margin: {section.left_margin}")
            logger.info(f"  - Right margin: {section.right_margin}")
        
        logger.info(f"Final page count: {page_breaks} (1 base page + {paragraph_breaks} paragraph breaks + {table_breaks} table breaks)")
        return page_breaks
    except Exception as e:
        logger.error(f"Error getting page count: {e}", exc_info=True)
        return 1

def test_template(template_path, data_path, output_path):
    """Тестирует шаблон и подсчет страниц."""
    try:
        # Загружаем шаблон
        logger.info(f"Loading template from: {template_path}")
        template = DocxTemplate(template_path)
        
        # Читаем тестовые данные
        logger.info(f"Reading test data from: {data_path}")
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Устанавливаем начальное значение страниц
        data['num_pages'] = 0
        
        # Рендерим документ первый раз
        logger.info("First render with num_pages = 0")
        template.render(data)
        template.save(output_path)
        
        # Считаем страницы
        pages = get_docx_pages(output_path)
        logger.info(f"First count: {pages} pages")
        
        # Обновляем количество страниц и рендерим снова
        data['num_pages'] = pages - 1
        logger.info(f"Second render with num_pages = {pages-1}")
        template.render(data)
        template.save(output_path)
        
        # Проверяем финальный результат
        final_pages = get_docx_pages(output_path)
        logger.info(f"Final count: {final_pages} pages")
        
        return True
    except Exception as e:
        logger.error(f"Test failed: {e}", exc_info=True)
        return False

def main():
    if len(sys.argv) != 4:
        print("Usage: test_pages.py <template_path> <data_path> <output_path>")
        sys.exit(1)

    template_path = sys.argv[1]
    data_path = sys.argv[2]
    output_path = sys.argv[3]

    if not test_template(template_path, data_path, output_path):
        sys.exit(1)

if __name__ == "__main__":
    main() 